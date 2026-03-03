#!/usr/bin/env python3
"""Generate index.html from the Hutt SCBU Handbook .docx file."""

import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import html as html_mod

doc = Document("/Users/michaelhewsonair/Desktop/Hutt/Hutt SCBU Handbook Jan 2026.docx")

# ── Helpers ──────────────────────────────────────────────────────────────────

def esc(text):
    return html_mod.escape(text)

def slug(text):
    s = re.sub(r'[^\w\s-]', '', text.lower())
    return re.sub(r'[\s_]+', '-', s).strip('-')[:60]

def para_to_html(para):
    """Convert a paragraph's runs into inline HTML (bold, italic, underline, superscript, subscript)."""
    if not para.runs:
        return esc(para.text)
    parts = []
    for run in para.runs:
        t = esc(run.text)
        if not t:
            continue
        if run.font.superscript:
            t = f'<sup>{t}</sup>'
        if run.font.subscript:
            t = f'<sub>{t}</sub>'
        if run.bold:
            t = f'<strong>{t}</strong>'
        if run.italic:
            t = f'<em>{t}</em>'
        if run.underline:
            t = f'<u>{t}</u>'
        parts.append(t)
    return ''.join(parts)

def linkify(text_html):
    """Turn URLs in already-escaped HTML into <a> links."""
    return re.sub(
        r'(https?://[^\s<>&]+)',
        r'<a href="\1" target="_blank" rel="noopener">\1</a>',
        text_html
    )

def table_to_html(table):
    """Convert a docx table to HTML."""
    rows_html = []
    for i, row in enumerate(table.rows):
        tag = 'th' if i == 0 else 'td'
        cells = []
        for cell in row.cells:
            content = esc(cell.text).replace('\n', '<br>')
            cells.append(f'<{tag}>{content}</{tag}>')
        rows_html.append('<tr>' + ''.join(cells) + '</tr>')
    header = '<thead>' + rows_html[0] + '</thead>' if rows_html else ''
    body = '<tbody>' + '\n'.join(rows_html[1:]) + '</tbody>' if len(rows_html) > 1 else ''
    return f'<div class="table-wrap"><table>{header}\n{body}</table></div>'


# ── Walk document elements in order (paragraphs AND tables interleaved) ──────

from docx.oxml.ns import qn

def iter_block_items(parent):
    """Yield paragraphs and tables in document order."""
    body = parent.element.body
    for child in body:
        if child.tag == qn('w:p'):
            yield doc.paragraphs[list(body.iterchildren(qn('w:p'))).index(child)]
        elif child.tag == qn('w:tbl'):
            yield doc.tables[list(body.iterchildren(qn('w:tbl'))).index(child)]

# Precompute paragraph index -> table mapping isn't needed with iter_block_items

# ── Define sections for the sidebar nav ──────────────────────────────────────

# Major section definitions: (nav_title, heading_match_text, subsections[])
SECTIONS = [
    ("Quick Links", "Quick links:", []),
    ("SCBU Introduction", "SCBU – Introduction", [
        ("Admission Criteria", "Admission criteria:"),
        ("Discharge", "Discharge:"),
        ("Daily Duties", "Daily duties on SCBU:"),
        ("Attending Deliveries", "Attending deliveries:"),
        ("Newborn Examination", "Newborn physical examination:"),
        ("Bloods", "Bloods:"),
        ("Investigations & Screening", "Other investigations and screening"),
        ("Referrals", "Referrals for babies in SCBU:"),
    ]),
    ("Neonatal Resuscitation", "Neonatal Resuscitation", []),
    ("Intubation Drugs", "SCBU Intubation Drugs Reference Chart", []),
    ("Fluids & Nutrition", "Fluids and Nutrition", [
        ("Feed Fortification", "Feed fortification"),
        ("Caffeine", "Caffeine"),
    ]),
    ("Respiratory", "Transient Tachypnoea", [
        ("TTN", "Transient Tachypnoea"),
        ("RDS", "Respiratory Distress Syndrome"),
        ("MAS", "Meconium Aspiration"),
        ("PPHN", "Persistent pulmonary hypertension"),
    ]),
    ("Electrolytes", "Electrolyte imbalance", [
        ("Hyponatraemia", "Hyponatraemia"),
        ("Hypernatraemia", "Hypernatraemia"),
        ("Hypokalaemia", "Hypokalaemia"),
        ("Hyperkalaemia", "Hyperkalaemia"),
    ]),
    ("Cardiac", "Cardiac issues", [
        ("Bradycardia", "Bradycardia during sleep"),
        ("CHD Screening", "Congenital heart disease screening"),
        ("Murmurs", "Cardiac murmurs"),
        ("ECG", "Neonatal ECG"),
        ("Congenital Heart Defect", "Congenital heart defect"),
        ("SVT", "Supraventricular Tachycardia"),
    ]),
    ("Birth Injuries", "Cephalohaematoma", []),
    ("Jaundice", "Jaundice", []),
    ("Haematology", "Haematology", [
        ("HDN", "Haemolytic disease"),
        ("Transfusion", "Transfusion Threshold"),
        ("Polycythaemia", "Polycythaemia"),
        ("Thrombocytopaenia", "Thrombocytopaenia"),
        ("NAIT", "NAIT"),
        ("Neutropaenia", "Neutropaenia"),
    ]),
    ("Infections", "Congenital and neonatal infections", [
        ("Neonatal Sepsis", "Neonatal Sepsis"),
        ("Asymptomatic Babies", "Management of asymptomatic"),
        ("Symptomatic Babies", "MANAGEMENT OF SYMPTOMATIC"),
    ]),
    ("Neurology", "Neurology", [
        ("HIE", "Hypoxic Ischaemic"),
        ("Seizures", "Neonatal Seizures"),
    ]),
    ("Metabolic / Endocrine", "Metabolic / Endocrine", [
        ("Hypoglycaemia", "Hypoglycaemia"),
        ("Inborn Errors", "Inborn errors"),
        ("Metabolic Bone Disease", "Metabolic Bone Disease"),
        ("Thyroid Disease", "Maternal thyroid"),
        ("DSD", "Ambiguous Genitalia"),
    ]),
    ("Gastrointestinal", "Gastrointestinal", [
        ("Bilious Vomiting", "Bilious vomiting"),
        ("NEC", "Necrotising enterocolitis"),
    ]),
    ("Skin Conditions", "Skin conditions", [
        ("Staph Infections", "Staphylococcal"),
        ("Herpes Simplex", "Herpes simplex"),
        ("Benign Skin", "Benign skin"),
        ("Sticky Eye", "Sticky eye"),
        ("Sacral Dimples", "Lumbar skin stigmata"),
    ]),
    ("Surgical Issues", "Surgical issues", [
        ("Undescended Testes", "Undescended Testes"),
        ("Hypospadias", "Hypospadias"),
        ("Inguinal Hernias", "Inguinal hernias"),
        ("Renal Tract", "Renal Tract"),
    ]),
    ("Resources", "Resources", []),
]

# ── Build content from document ──────────────────────────────────────────────

# Walk through document and build HTML sections
content_blocks = []  # list of (heading_text, level, html_content)
current_html = []
current_heading = "preamble"
current_level = 0
skip_toc = True  # skip until we hit first real Heading 2

# Track list state
in_list = False
list_depth = 0

def flush_list():
    global in_list, list_depth
    if in_list:
        current_html.append('</ul>' * max(list_depth, 1))
        in_list = False
        list_depth = 0

def start_new_section(heading_text, level):
    global current_html, current_heading, current_level, skip_toc
    flush_list()
    if current_html and current_heading != "preamble":
        content_blocks.append((current_heading, current_level, '\n'.join(current_html)))
    current_html = []
    current_heading = heading_text
    current_level = level
    skip_toc = False

for block in iter_block_items(doc):
    # Table
    if hasattr(block, 'rows'):
        if skip_toc:
            continue
        flush_list()
        # Skip empty/placeholder tables
        total_text = ''.join(cell.text for row in block.rows for cell in row.cells).strip()
        if total_text:
            current_html.append(table_to_html(block))
        continue

    # Paragraph
    para = block
    style = para.style.name
    text = para.text.strip()

    # Skip TOC entries and empty paragraphs in preamble
    if 'TOC' in style:
        continue
    if not text:
        continue

    # Detect headings
    is_heading = False
    if style.startswith('Heading'):
        if text:  # skip empty headings
            level = 2
            if '3' in style:
                level = 3
            elif style == 'Heading':
                level = 2
            start_new_section(text, level)
            sid = slug(text)
            tag = f'h{level}'
            current_html.append(f'<{tag} id="{sid}">{linkify(para_to_html(para))}</{tag}>')
            is_heading = True

    if is_heading or skip_toc:
        continue

    # Detect bold standalone lines as sub-headings (h3/h4)
    all_bold = para.runs and all(r.bold for r in para.runs if r.text.strip())
    is_list_style = 'List' in style

    if all_bold and not is_list_style and len(text) < 120 and not text.startswith('•'):
        flush_list()
        sid = slug(text)
        current_html.append(f'<h4 id="{sid}">{linkify(para_to_html(para))}</h4>')
        continue

    # List items
    if is_list_style or text.startswith('•') or text.startswith('-'):
        if not in_list:
            current_html.append('<ul>')
            in_list = True
            list_depth = 1
        content = linkify(para_to_html(para))
        # Remove leading bullet char
        content = re.sub(r'^[•\-]\s*', '', content)
        if all_bold and len(text) < 100:
            current_html.append(f'<li><strong>{content}</strong></li>')
        else:
            current_html.append(f'<li>{content}</li>')
        continue

    # Regular paragraph
    flush_list()
    content = linkify(para_to_html(para))
    if all_bold:
        current_html.append(f'<p><strong>{content}</strong></p>')
    else:
        current_html.append(f'<p>{content}</p>')

# Flush last section
flush_list()
if current_html and current_heading != "preamble":
    content_blocks.append((current_heading, current_level, '\n'.join(current_html)))

# ── Build navigation sidebar HTML ───────────────────────────────────────────

# Create a mapping from heading text -> slug for nav links
heading_slugs = {}
for heading_text, level, html_content in content_blocks:
    heading_slugs[heading_text] = slug(heading_text)

def find_slug(match_text):
    """Find slug for a heading that starts with match_text."""
    for h, s in heading_slugs.items():
        if match_text.lower() in h.lower():
            return s
    return slug(match_text)

nav_html = '<ul class="nav-list">\n'
for sec_title, match_text, subsections in SECTIONS:
    sec_slug = find_slug(match_text)
    if subsections:
        nav_html += f'  <li class="nav-item has-sub">\n'
        nav_html += f'    <a href="#{sec_slug}">{esc(sec_title)}</a>\n'
        nav_html += f'    <ul class="nav-sub">\n'
        for sub_title, sub_match in subsections:
            sub_slug = find_slug(sub_match)
            nav_html += f'      <li><a href="#{sub_slug}">{esc(sub_title)}</a></li>\n'
        nav_html += f'    </ul>\n'
        nav_html += f'  </li>\n'
    else:
        nav_html += f'  <li class="nav-item"><a href="#{sec_slug}">{esc(sec_title)}</a></li>\n'
nav_html += '</ul>\n'

# ── Assemble main content HTML ───────────────────────────────────────────────

# Group content blocks into major sections for <section> wrappers
main_html = ''
for heading_text, level, html_content in content_blocks:
    sid = slug(heading_text)
    main_html += f'<section id="sec-{sid}" class="content-section">\n{html_content}\n</section>\n\n'

# ── Write full HTML ─────────────────────────────────────────────────────────

full_html = f'''<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Hutt SCBU Handbook</title>
<style>
/* ── Reset & Base ────────────────────────────────────── */
*, *::before, *::after {{ box-sizing: border-box; margin: 0; padding: 0; }}
html {{ scroll-behavior: smooth; scroll-padding-top: 1rem; }}
body {{
  font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif;
  font-size: 15px;
  line-height: 1.6;
  color: #1a1a1a;
  background: #f5f6f8;
}}

/* ── Sidebar ─────────────────────────────────────────── */
.sidebar {{
  position: fixed;
  top: 0; left: 0;
  width: 280px;
  height: 100vh;
  background: #1e293b;
  color: #e2e8f0;
  overflow-y: auto;
  padding: 1rem 0;
  z-index: 100;
  transition: transform 0.3s ease;
}}
.sidebar-header {{
  padding: 0.75rem 1.25rem 1rem;
  border-bottom: 1px solid #334155;
}}
.sidebar-header h1 {{
  font-size: 1.1rem;
  font-weight: 700;
  color: #fff;
  line-height: 1.3;
}}
.sidebar-header p {{
  font-size: 0.75rem;
  color: #94a3b8;
  margin-top: 0.25rem;
}}
.nav-list {{
  list-style: none;
  padding: 0.5rem 0;
}}
.nav-item > a, .nav-sub a {{
  display: block;
  padding: 0.4rem 1.25rem;
  color: #cbd5e1;
  text-decoration: none;
  font-size: 0.85rem;
  border-left: 3px solid transparent;
  transition: all 0.15s;
}}
.nav-item > a:hover, .nav-sub a:hover {{
  background: #334155;
  color: #fff;
  border-left-color: #60a5fa;
}}
.nav-item > a.active {{
  background: #334155;
  color: #60a5fa;
  border-left-color: #60a5fa;
  font-weight: 600;
}}
.nav-sub {{
  list-style: none;
  padding: 0;
  display: none;
}}
.nav-item.has-sub.open > .nav-sub {{
  display: block;
}}
.nav-sub a {{
  padding-left: 2rem;
  font-size: 0.8rem;
  color: #94a3b8;
}}
.nav-sub a:hover {{
  color: #e2e8f0;
}}

/* ── Hamburger ───────────────────────────────────────── */
.menu-toggle {{
  display: none;
  position: fixed;
  top: 0.75rem; left: 0.75rem;
  z-index: 200;
  background: #1e293b;
  color: #fff;
  border: none;
  border-radius: 6px;
  width: 40px; height: 40px;
  font-size: 1.3rem;
  cursor: pointer;
  line-height: 1;
}}
.overlay {{
  display: none;
  position: fixed;
  inset: 0;
  background: rgba(0,0,0,0.4);
  z-index: 50;
}}

/* ── Main content ────────────────────────────────────── */
.main {{
  margin-left: 280px;
  padding: 2rem 2.5rem 4rem;
  max-width: 900px;
}}
.content-section {{
  background: #fff;
  border-radius: 8px;
  padding: 1.5rem 2rem;
  margin-bottom: 1.5rem;
  box-shadow: 0 1px 3px rgba(0,0,0,0.06);
}}

/* ── Typography ──────────────────────────────────────── */
h2 {{
  font-size: 1.5rem;
  color: #1e293b;
  border-bottom: 2px solid #e2e8f0;
  padding-bottom: 0.5rem;
  margin-bottom: 1rem;
}}
h3 {{
  font-size: 1.2rem;
  color: #334155;
  margin: 1.5rem 0 0.75rem;
}}
h4 {{
  font-size: 1rem;
  color: #475569;
  margin: 1.25rem 0 0.5rem;
}}
p {{
  margin-bottom: 0.6rem;
}}
a {{
  color: #2563eb;
  text-decoration: none;
}}
a:hover {{
  text-decoration: underline;
}}
ul, ol {{
  margin: 0.5rem 0 0.75rem 1.5rem;
}}
li {{
  margin-bottom: 0.3rem;
}}

/* ── Tables ──────────────────────────────────────────── */
.table-wrap {{
  overflow-x: auto;
  margin: 1rem 0;
}}
table {{
  width: 100%;
  border-collapse: collapse;
  font-size: 0.85rem;
}}
th, td {{
  border: 1px solid #e2e8f0;
  padding: 0.5rem 0.75rem;
  text-align: left;
  vertical-align: top;
}}
th {{
  background: #f1f5f9;
  font-weight: 600;
  color: #334155;
  white-space: nowrap;
}}
tbody tr:nth-child(even) {{
  background: #f8fafc;
}}

/* ── Search box ──────────────────────────────────────── */
.search-box {{
  padding: 0.75rem 1.25rem;
}}
.search-box input {{
  width: 100%;
  padding: 0.4rem 0.6rem;
  border-radius: 4px;
  border: 1px solid #475569;
  background: #334155;
  color: #e2e8f0;
  font-size: 0.8rem;
  outline: none;
}}
.search-box input::placeholder {{ color: #94a3b8; }}
.search-box input:focus {{ border-color: #60a5fa; }}

/* ── Mobile ──────────────────────────────────────────── */
@media (max-width: 768px) {{
  .sidebar {{
    transform: translateX(-100%);
  }}
  .sidebar.open {{
    transform: translateX(0);
  }}
  .overlay.open {{
    display: block;
  }}
  .menu-toggle {{
    display: block;
  }}
  .main {{
    margin-left: 0;
    padding: 3.5rem 1rem 3rem;
  }}
  .content-section {{
    padding: 1rem 1.25rem;
  }}
}}

/* ── Print ───────────────────────────────────────────── */
@media print {{
  .sidebar, .menu-toggle, .overlay, .search-box {{
    display: none !important;
  }}
  .main {{
    margin-left: 0;
    padding: 0;
    max-width: 100%;
  }}
  .content-section {{
    box-shadow: none;
    border: none;
    page-break-inside: avoid;
    padding: 0.5rem 0;
  }}
  h2 {{ page-break-after: avoid; }}
  table {{ page-break-inside: avoid; }}
}}
</style>
</head>
<body>

<button class="menu-toggle" aria-label="Toggle menu">&#9776;</button>
<div class="overlay"></div>

<nav class="sidebar">
  <div class="sidebar-header">
    <h1>Hutt SCBU Handbook</h1>
    <p>Quick Reference Guide &middot; 2026</p>
  </div>
  <div class="search-box">
    <input type="text" id="nav-search" placeholder="Search sections&hellip;">
  </div>
  {nav_html}
</nav>

<main class="main">
  <div class="preamble content-section">
    <h1 style="font-size:1.6rem; margin-bottom:0.5rem">Hutt SCBU Handbook</h1>
    <p style="color:#64748b; font-size:0.9rem">A Quick Reference Guide of Practice and Common Problems</p>
    <p style="color:#64748b; font-size:0.8rem; margin-top:0.5rem">
      This document is for local use only in the SCBU at Hutt Hospital.<br>
      It reflects practice recommendations for junior doctors in this unit.<br>
      Latest update: Nov 2024
    </p>
  </div>
  {main_html}
</main>

<script>
// Sidebar toggle (mobile)
const toggle = document.querySelector('.menu-toggle');
const sidebar = document.querySelector('.sidebar');
const overlay = document.querySelector('.overlay');
toggle.addEventListener('click', () => {{
  sidebar.classList.toggle('open');
  overlay.classList.toggle('open');
}});
overlay.addEventListener('click', () => {{
  sidebar.classList.remove('open');
  overlay.classList.remove('open');
}});

// Collapsible nav sections
document.querySelectorAll('.nav-item.has-sub > a').forEach(link => {{
  link.addEventListener('click', (e) => {{
    // Allow navigation but toggle sub-menu
    link.parentElement.classList.toggle('open');
  }});
}});

// Active section highlight on scroll
const navLinks = document.querySelectorAll('.nav-list a');
const sections = document.querySelectorAll('.content-section');
const observer = new IntersectionObserver(entries => {{
  entries.forEach(entry => {{
    if (entry.isIntersecting) {{
      const id = entry.target.querySelector('h2, h3')?.id;
      if (!id) return;
      navLinks.forEach(l => l.classList.remove('active'));
      const active = document.querySelector(`.nav-list a[href="#${{id}}"]`);
      if (active) {{
        active.classList.add('active');
        // Open parent submenu
        const parent = active.closest('.has-sub');
        if (parent) parent.classList.add('open');
      }}
    }}
  }});
}}, {{ rootMargin: '-20% 0px -70% 0px' }});
sections.forEach(s => observer.observe(s));

// Close sidebar on nav click (mobile)
navLinks.forEach(link => {{
  link.addEventListener('click', () => {{
    if (window.innerWidth <= 768) {{
      sidebar.classList.remove('open');
      overlay.classList.remove('open');
    }}
  }});
}});

// Search filter
document.getElementById('nav-search').addEventListener('input', function() {{
  const q = this.value.toLowerCase();
  document.querySelectorAll('.nav-item').forEach(item => {{
    const text = item.textContent.toLowerCase();
    item.style.display = text.includes(q) ? '' : 'none';
  }});
}});
</script>
</body>
</html>'''

with open('/Users/michaelhewsonair/Desktop/Hutt/index.html', 'w') as f:
    f.write(full_html)

print(f"Generated index.html ({len(full_html):,} bytes)")
print(f"Content sections: {len(content_blocks)}")
for h, l, c in content_blocks:
    print(f"  H{l}: {h[:60]}")
